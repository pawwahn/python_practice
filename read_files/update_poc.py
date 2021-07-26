from docx import Document
from datetime import date

# open the document
doc = Document('D:/Lv/poc.docx')

sym = "£ "
cur_date = date.today()

quote_line = 'Attached is an illustration showing how the above changes will effect the plan, ' \
             'by taking this payment the Money Purchase Annual Allowance (MPAA) rules have been triggered.'

risk_line = 'Please find attached a copy of your Risk Warnings based upon the survey completed. ' \
            'Please consider carefully how the money being withdrawn is used, and read the attached ‘Pension Scam’ flyer, ' \
            'on the steps that can be taken to protect the money in the pension.'

table1 = {"gross_amount": sym + "10000", "tax_amount": sym + "100", "net_amount": sym + "9000"}
table2 = {"member_name": "040447", "plan_number": "p040447", "adviser_charge": sym + "100"}

adv_data = {"oprn": "SdX 12", 'account_name': 'Gen Mark', "adviser_name": "Bajaj Finserver",
            "adv_add_line1": 'Hotel Turya', "adv_add_line2": '11 Floor', 'adv_add_line3': 'Sholinganallur',
            'DATE': str(cur_date), 'quote': False,'risk_warnings': False, 'adviser_charges_paid': False}


for i in adv_data:
    for p in doc.paragraphs:
        if adv_data['quote'] == False :
            p.text = p.text.replace(quote_line, " ")
        if adv_data['risk_warnings'] == False :
            p.text = p.text.replace(risk_line, " ")
        if adv_data['adviser_charges_paid'] == True :
            p.text = p.text.replace('adviser_name',adv_data['adviser_name'])
            p.text = p.text.replace('adv_add_line1',adv_data['adv_add_line1'])
            p.text = p.text.replace('adv_add_line2',adv_data['adv_add_line2'])
            p.text = p.text.replace('adv_add_line3',adv_data['adv_add_line3'])
        else:
            p.text = p.text.replace('adviser_name', " ")
            p.text = p.text.replace('adv_add_line1', " ")
            p.text = p.text.replace('adv_add_line2', " ")
            p.text = p.text.replace('adv_add_line3', " ")
        if p.text.find(i) >= 0:
            p.text = p.text.replace(i, adv_data[i])

doc.tables
# tables 1
# print("Retrieved value: " + doc.tables[0].cell(0, 0).text)
doc.tables[0].cell(0, 1).text = table1['gross_amount']
doc.tables[0].cell(1, 1).text = table1['tax_amount']
doc.tables[0].cell(2, 1).text = table1['net_amount']

# table2
doc.tables[1].cell(1, 0).text = table2['member_name'].center(40)
doc.tables[1].cell(1, 1).text = table2['plan_number'].center(45)
doc.tables[1].cell(1, 2).text = table2['adviser_charge'].center(45)

# save changed document
doc.save('D:/Lv/updated_poc.docx')

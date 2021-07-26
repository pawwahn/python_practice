import docx,os
from docx.shared import Inches, RGBColor, Pt
try:
    os.remove('C:/Users/56007/PycharmProjects/python_practice/read_files/lv.docx')
except:
    pass
doc = docx.Document()

section = doc.sections[0]
footer = section.footer
footer_para1 = footer.paragraphs[0]
# footer_para1.font.size = Pt(12)
footer_para1.text = "LV=, Savings & Retirement, Pease House, Tilehouse Street, Hitchin, Herts, SG5 2DX \n" \
                    "LV= and Liverpool Victoria are registered trademarks of Liverpool Victoria Financial Services Limited and LV= and LV= Liverpool Victoria are trading styles of the Liverpool Victoria group of companies. Liverpool Victoria Financial Services Limited, registered in England with registration number 12383237 is authorised by the Prudential Regulation Authority and regulated by the Financial Conduct Authority and the Prudential Regulation Authority, register number 110035. NM Pensions Trustees Limited, (registered in England No. 4299742), acts as scheme trustee. Registered address for both companies: County Gates, Bournemouth, BH1 2NF. Phone: 01202 292333.32453-2020 05/20"


doc.add_paragraph("\t\t\t\t\t")
doc.add_picture('C:/Users/56007/PycharmProjects/python_practice/read_files/download.jfif',
                width=Inches(0.6),height=Inches(0.6))
# p1 = doc.add_paragraph()
para1 = doc.add_paragraph("Hello,\n \n")
para1.add_run("Thank you for your instructions, we've now processed the income request. \n \n")
para1.add_run("We've made the following payment today, it should show in the target account within 3-5 working days; ")

# Table data in a form of list
gp, td, np = 100, 200, 300
data = (('Gross Payment', str(gp)), ('Tax Deducted', str(td)), ('Net Payment', str(np)))
# Creating a table object
table = doc.add_table(rows=1, cols=2)

# Adding heading in the 1st row of the table
row = table.rows[0].cells
row[0].text = 'Details'
row[1].text = 'Amount'

# Adding data from the list to the table
for id, name in data:
    # Adding a row and then adding data in it.
    row = table.add_row().cells
    row[0].text = str(id)
    row[1].text = name

# Adding style to a table
table.style = 'Colorful List'

v1 = 200

para2 = doc.add_paragraph("\n")
para2.add_run("The income will be subject to tax, and the tax code currently in operation is {} \n\n".format(v1))
para2.add_run("If you have any questions concerning the tax or tax code, please refer to the "
              "attached ‘Tax on Income Lump Sums’ document for useful information. "
              "Alternatively, contact HM Revenue and Customs on 0300 200 3300, or via their website ")
para2.add_run("www.hmrc.gov.uk.")
para2.add_run(" They may ask for our PAYE reference which is 321/XA13246. \n \n")

para2.add_run("We are able to set up a regular income from your plan which can be paid Monthly, "
              "Annually, Half Yearly and Quarterly. We may also be able to set up a regular sale"
              " of funds to cover the income payments but we will inform you if we require an "
              "instruction to place funds on your account to cover future payments. "
              "If you would like to take advantage of this option, please do not hesitate to contact us "
              "and we will talk you through what is required \n\n")

# font_clr_red = para2.add_run('(If Quote Issued keep the below statement/If no Quote issued remove it)')
# font_clr_red.bold=True
#
# # Adding forest green colour to the text
# # RGBColor(R, G, B)
# font_clr_red.font.color.rgb = RGBColor(255, 0, 0)

# If quote Issued -- condition to be added
para2.add_run("Attached is an illustration showing how the above changes will effect the plan, by taking this payment the ")
para2.add_run("Money Purchase Annual Allowance (MPAA) ").bold=True
para2.add_run("rules have been triggered. \n\n")

# If Risk Warnings
para2.add_run("Please find attached a copy of your Risk Warnings based upon the survey completed. Please consider carefully how the money being withdrawn is used, and read the attached ‘Pension Scam’ flyer, on the steps that can be taken to protect the money in the pension. \n\n")
para2.add_run("If you have any questions, please contact us and we’ll be happy to help! \n")

blue_cnt = doc.add_heading("Your LV= Drawdown Team", 1)
para2.add_run("\n\n")
# blue_cnt.bold = True
# blue_cnt.font.color.rgb= RGBColor(0, 136, 204)











doc.save('C:/Users/56007/PycharmProjects/python_practice/read_files/lv.docx')

from datetime import date
from docx import Document

# This Python file uses the following encoding: utf-8

# open the document
doc = Document('D:/Lv/poc.docx')

sym = "£ "
cur_date = date.today()

quote_line = 'Attached is an illustration showing how the above changes will effect the plan, ' \
             'by taking this payment the Money Purchase Annual Allowance (MPAA) rules have been triggered.'

risk_line = "Please find attached a copy of your Risk Warnings based upon the survey completed. " \
            "Please consider carefully how the money being withdrawn is used, and read the attached 'Pension Scam' flyer, " \
            "on the steps that can be taken to protect the money in the pension."

table1 = {"gross_amount": sym + "10000", "tax_amount": sym + "100", "net_amount": sym + "9000"}
table2 = {"member_name": "040447", "plan_number": "p040447", "adviser_charge": sym + "100"}

adv_data = {"operation": "Operation 1", 'account_name': 'Gen Mark', "adviser_name": "Bajaj Finserver",
            "adv_add_line1": 'Hotel Turya', "adv_add_line2": '11 Floor', 'adv_add_line3': 'Sholinganallur',
            'DATE': str(cur_date), 'quote': True, 'risk_warnings': True, 'adviser_charges_paid': True}


def find_replace(paragraph_keyword, draft_keyword, paragraph):
    if paragraph_keyword in paragraph.text:
        paragraph.text = paragraph.text.replace(paragraph_keyword, draft_keyword)

for paragraph in doc.paragraphs:
    if not adv_data['quote']:
        find_replace(quote_line, "\n", paragraph)
    if not adv_data['risk_warnings']:
        find_replace(risk_line, "\n", paragraph)
    if not adv_data['adviser_charges_paid']:
        find_replace('adviser_name', "\n", paragraph)
        find_replace('adv_add_line1', "\n", paragraph)
        find_replace('adv_add_line2', "\n", paragraph)
        find_replace('adv_add_line3', "\n", paragraph)
        find_replace('DATE', "\n", paragraph)

for i in adv_data:
    for p in doc.paragraphs:
        if p.text.find(i) >= 0:
            p.text = p.text.replace(i, adv_data[i])



doc.tables
print(doc.tables)
# tables 1
# print("Retrieved value: " + doc.tables[0].cell(0, 0).text)
doc.tables[0].cell(0, 1).text = table1['gross_amount']
doc.tables[0].cell(1, 1).text = table1['tax_amount']
doc.tables[0].cell(2, 1).text = table1['net_amount']
#
# # table2
doc.tables[1].cell(1, 0).text = table2['member_name'].center(40)
doc.tables[1].cell(1, 1).text = table2['plan_number'].center(45)
doc.tables[1].cell(1, 2).text = table2['adviser_charge'].center(45)

# save changed document
doc.save('D:/Lv/updated_poc1.docx')

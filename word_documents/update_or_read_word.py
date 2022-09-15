import docx
from docx.shared import Cm, Pt

file_path = r'C:\Users\56007\Documents\testing_word.docx'

doc = docx.Document(file_path)
print(doc)
#print(doc.paragraphs)

for para in doc.paragraphs:
    print(para.text)

print(doc.tables)
#print(doc.tables[0].cell(0,1).tables[0].cell(0,0).text)

# customizing the table
table = doc.add_table(0, 0)  # we add rows iteratively
table.style = 'TableGrid'
first_column_width = 5
second_column_with = 10
table.add_column(Cm(first_column_width))
table.add_column(Cm(second_column_with))

for index, stat_item in enumerate(text_stats.items()):
    table.add_row()
    stat_name, stat_result = stat_item
    row = table.rows[index]
    row.cells[0].text = str(stat_name)
    row.cells[1].text = str(stat_result)
word_document.add_page_break()

doc.add_table(2,2,style=None)
table = doc._body.add_table(rows=2, cols=2,width=3)
#for tab in doc.tables:
#    print(tab)

doc.save(file_path)
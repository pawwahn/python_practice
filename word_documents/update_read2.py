import docx

file_path = r'C:\Users\56007\Documents\testing_word.docx'

doc = docx.Document(file_path)
print(doc)

for para in doc.paragraphs:
    print(para.text)

print(doc.tables)
print(doc.tables[0].cell(0,0).tables[0].cell(0,0).text)
print(doc.tables[0].cell(0,0).tables[0].cell(0,1).text)
#for tab in doc.tables:
#    print(tab)
import docx
import os

doc = docx.Document()

#doc.add_heading("Table Documemt", 0)

sipp_data = [
	[['data a', 'data b'],['data 1', 'data 2'],['data A', 'data B']],[['a','b'],['c','d'],['e','f']],
	[['data a', 'data b'],['data 1', 'data 2'],['data A', 'data B']],
	[['data a', 'data b'],['data 1', 'data 2'],['data A', 'data B']],
	[['data a', 'data b'],['data 1', 'data 2'],['data A', 'data B']],
	[['data a', 'data b'],['data 1', 'data 2'],['data A', 'data B']],
	[['data a', 'data b'],['data 1', 'data 2'],['data A', 'data B']],
	[['data a', 'data b'],['data 1', 'data 2'],['data A', 'data B']]
]

# mytab = doc.add_table(rows=2, cols=2)
# print(help(docx.table))
#print(mytab.rows[0])
# mytab.rows[0].add_table(rows=2,cols=2)
# mytab.style = 'Table Grid'

#mytab.tables[0].cell(0,1).tables[0].cell(0,0).text = 'pavan'

#print(doc.tables)

# doc.tables[0].add_table(rows=1,cols=2)
#
# #working code from
for sipp in sipp_data:
	#doc.add_heading("Table Documemt", 0)
	mytable = doc.add_table(rows=1, cols=2)
	mytable.style = 'Table Grid'
	dr_cells = mytable.rows[0].cells
	dr_cells[0].text = 'ID'
	dr_cells[1].text = 'Notes'
	for a,b in sipp:
		row_cells = mytable.add_row().cells
		row_cells[0].text = str(a)
		row_cells[1].text = str(b)
	para = doc.add_paragraph('\n')
#
# #working code to
#
# #doc.tables[0].cell(0,0).
#
#
file_path = r'C:\Users\56007\Documents\create_table.docx'
doc.save(file_path)
#
